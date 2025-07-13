import docx

text = docx.Document("file_name.docx")
# print(text.paragraphs[1].text)
data = {}
paras = text.paragraphs
# print(paras)
# print(len(paras))
for i in range(0, len(paras)):
    print(text.paragraphs[i].text)
    data[i] = tuple(paras[i].text.split('\t'))

print(data)
data_values = list(data.values())
print(data_values)